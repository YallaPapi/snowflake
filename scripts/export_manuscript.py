#!/usr/bin/env python3
"""
Manuscript Export Script for Auto-Novel Snowflake Engine
Converts markdown manuscript to DOCX and EPUB formats with professional formatting
"""

import os
import re
import sys
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import subprocess

def setup_document_styles(doc):
    """Set up professional manuscript styles"""
    
    # Check if styles already exist, if not create them
    style_names = [style.name for style in doc.styles]
    
    # Title style
    if 'Title Page' not in style_names:
        title_style = doc.styles.add_style('Title Page', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(24)
    
    # Author style
    if 'Author' not in style_names:
        author_style = doc.styles.add_style('Author', WD_STYLE_TYPE.PARAGRAPH)
        author_style.font.name = 'Times New Roman'
        author_style.font.size = Pt(16)
        author_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_style.paragraph_format.space_after = Pt(12)
    
    # Chapter title style
    if 'Chapter Title' not in style_names:
        chapter_style = doc.styles.add_style('Chapter Title', WD_STYLE_TYPE.PARAGRAPH)
        chapter_style.font.name = 'Times New Roman'
        chapter_style.font.size = Pt(18)
        chapter_style.font.bold = True
        chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chapter_style.paragraph_format.space_before = Pt(72)
        chapter_style.paragraph_format.space_after = Pt(24)
        chapter_style.paragraph_format.page_break_before = True
    
    # Body text style - use built-in Normal style and modify it
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 2.0  # Double spacing
    normal_style.paragraph_format.first_line_indent = Inches(0.5)
    normal_style.paragraph_format.space_after = Pt(0)
    
    return doc

def add_header_footer(doc, title, author):
    """Add header and footer with page numbers"""
    section = doc.sections[0]
    
    # Header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = f"{author} / {title.upper()}"
    header_para.style.font.name = 'Times New Roman'
    header_para.style.font.size = Pt(12)
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Footer with page numbers
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page number field
    run = footer_para.runs[0] if footer_para.runs else footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    run._element.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar2)

def parse_manuscript(content):
    """Parse markdown manuscript into structured content"""
    lines = content.split('\n')
    title = None
    author = None
    chapters = []
    current_chapter = None
    current_content = []
    
    for line in lines:
        line = line.strip()
        
        # Extract title (first occurrence of # title)
        if line.startswith('# ') and title is None:
            title = line[2:].strip()
            continue
            
        # Extract author (look for "By " line)
        if line.startswith('By ') and author is None:
            author = line[3:].strip()
            continue
            
        # Chapter headers
        if line.startswith('## Chapter ') or line.startswith('**Chapter '):
            # Save previous chapter
            if current_chapter is not None:
                chapters.append({
                    'title': current_chapter,
                    'content': '\n'.join(current_content).strip()
                })
            
            # Start new chapter
            if line.startswith('## Chapter '):
                current_chapter = line[3:].strip()
            else:
                current_chapter = line.replace('**', '').strip()
            current_content = []
            continue
            
        # Skip markdown separators and empty lines at start
        if line in ['---', ''] and not current_content:
            continue
            
        # Add content to current chapter
        if current_chapter is not None:
            current_content.append(line)
    
    # Add final chapter
    if current_chapter is not None:
        chapters.append({
            'title': current_chapter,
            'content': '\n'.join(current_content).strip()
        })
    
    return title, author, chapters

def create_docx(manuscript_path, output_path):
    """Create DOCX format manuscript"""
    print(f"Creating DOCX: {output_path}")
    
    # Read manuscript
    with open(manuscript_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Parse content
    title, author, chapters = parse_manuscript(content)
    
    if not title:
        title = "CODE OF DECEPTION"
    if not author:
        author = "The Snowflake Engine"
    
    # Create document
    doc = Document()
    doc = setup_document_styles(doc)
    
    # Set up page margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title page
    title_para = doc.add_paragraph(title, style='Title Page')
    subtitle_para = doc.add_paragraph("A Techno-Thriller Novel", style='Author')
    doc.add_paragraph()  # Empty line
    doc.add_paragraph()  # Empty line
    author_para = doc.add_paragraph(f"By {author}", style='Author')
    
    # Add page break after title page
    doc.add_page_break()
    
    # Add header and footer
    add_header_footer(doc, title, author)
    
    # Add chapters
    for i, chapter in enumerate(chapters):
        # Chapter title
        chapter_para = doc.add_paragraph(chapter['title'], style='Chapter Title')
        
        # Chapter content
        paragraphs = chapter['content'].split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text and not para_text.startswith('**Chapter'):
                # Clean up any remaining markdown
                para_text = re.sub(r'\*\*(.*?)\*\*', r'\1', para_text)  # Remove bold markdown
                para_text = re.sub(r'\*(.*?)\*', r'\1', para_text)      # Remove italic markdown
                
                doc.add_paragraph(para_text)
    
    # Save document
    doc.save(output_path)
    print(f"DOCX created successfully: {output_path}")

def create_epub(manuscript_path, output_path):
    """Create EPUB format using simple internal creator"""
    print(f"Creating EPUB: {output_path}")
    
    try:
        # Import our simple EPUB creator
        from simple_epub_creator import create_epub_structure
        
        # Read manuscript
        with open(manuscript_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Parse content
        title, author, chapters = parse_manuscript(content)
        
        if not title:
            title = "CODE OF DECEPTION"
        if not author:
            author = "The Snowflake Engine"
        
        # Create EPUB
        create_epub_structure(title, author, chapters, output_path)
        
    except Exception as e:
        print(f"EPUB creation failed: {e}")
        # Fallback to pandoc if available
        try:
            subprocess.run(['pandoc', '--version'], capture_output=True, check=True)
            
            cmd = [
                'pandoc',
                manuscript_path,
                '-o', output_path,
                '--toc',
                '--toc-depth=2'
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode == 0:
                print(f"EPUB created successfully with pandoc: {output_path}")
            else:
                print(f"Pandoc EPUB creation also failed: {result.stderr}")
                
        except (subprocess.CalledProcessError, FileNotFoundError):
            print("Pandoc not available either. Could not create EPUB.")

def main():
    """Main function"""
    if len(sys.argv) > 1:
        manuscript_path = sys.argv[1]
    else:
        # Default path
        manuscript_path = r"C:\Users\Stuart\Desktop\Projects\snowflake\artifacts\code_of_deception_20250821_212841\COMPLETE_MANUSCRIPT.md"
    
    if not os.path.exists(manuscript_path):
        print(f"Error: Manuscript file not found: {manuscript_path}")
        sys.exit(1)
    
    # Set up output directory
    manuscript_dir = Path(manuscript_path).parent
    base_name = Path(manuscript_path).stem
    
    # Output paths
    docx_path = manuscript_dir / f"{base_name}.docx"
    epub_path = manuscript_dir / f"{base_name}.epub"
    
    print(f"Processing manuscript: {manuscript_path}")
    print(f"Word count: {subprocess.run(['wc', '-w', manuscript_path], capture_output=True, text=True).stdout.strip()}")
    
    # Create DOCX
    try:
        create_docx(manuscript_path, docx_path)
    except Exception as e:
        print(f"Error creating DOCX: {e}")
    
    # Create EPUB
    try:
        create_epub(manuscript_path, epub_path)
    except Exception as e:
        print(f"Error creating EPUB: {e}")
    
    print("\nExport complete!")
    print(f"Files created in: {manuscript_dir}")

if __name__ == "__main__":
    main()