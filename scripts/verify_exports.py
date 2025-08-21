#!/usr/bin/env python3
"""
Export Verification Script for Auto-Novel Snowflake Engine
Validates exported DOCX and EPUB files
"""

import os
import sys
import zipfile
from pathlib import Path
from docx import Document

def verify_docx(docx_path):
    """Verify DOCX file structure and content"""
    print(f"\nVerifying DOCX: {docx_path}")
    
    try:
        doc = Document(docx_path)
        
        # Count paragraphs and estimated words
        paragraph_count = len(doc.paragraphs)
        word_count = 0
        
        for para in doc.paragraphs:
            word_count += len(para.text.split())
        
        print(f"  [OK] File opens successfully")
        print(f"  [OK] Paragraphs: {paragraph_count}")
        print(f"  [OK] Estimated words: {word_count:,}")
        
        # Check for title and chapter structure
        has_title = False
        chapter_count = 0
        
        for para in doc.paragraphs[:10]:  # Check first 10 paragraphs
            text = para.text.strip()
            if "CODE OF DECEPTION" in text:
                has_title = True
                break
        
        for para in doc.paragraphs:
            if "Chapter" in para.text and len(para.text.strip()) < 50:
                chapter_count += 1
        
        print(f"  [OK] Title page: {'Found' if has_title else 'Not found'}")
        print(f"  [OK] Chapters detected: {chapter_count}")
        
        # Check file size
        file_size = os.path.getsize(docx_path)
        print(f"  [OK] File size: {file_size:,} bytes ({file_size/1024:.1f} KB)")
        
        return True
        
    except Exception as e:
        print(f"  [ERROR] Error: {e}")
        return False

def verify_epub(epub_path):
    """Verify EPUB file structure"""
    print(f"\nVerifying EPUB: {epub_path}")
    
    try:
        with zipfile.ZipFile(epub_path, 'r') as epub:
            file_list = epub.namelist()
            
            # Check required EPUB structure
            required_files = ['mimetype', 'META-INF/container.xml']
            optional_files = ['OEBPS/content.opf', 'OEBPS/toc.ncx']
            
            print(f"  [OK] Total files in EPUB: {len(file_list)}")
            
            for req_file in required_files:
                if req_file in file_list:
                    print(f"  [OK] Required file found: {req_file}")
                else:
                    print(f"  [ERROR] Missing required file: {req_file}")
                    return False
            
            for opt_file in optional_files:
                if opt_file in file_list:
                    print(f"  [OK] Optional file found: {opt_file}")
            
            # Check for chapter files
            chapter_files = [f for f in file_list if f.startswith('OEBPS/chapter') and f.endswith('.xhtml')]
            print(f"  [OK] Chapter files found: {len(chapter_files)}")
            
            # Check mimetype
            mimetype = epub.read('mimetype').decode('utf-8').strip()
            if mimetype == 'application/epub+zip':
                print(f"  [OK] Correct mimetype: {mimetype}")
            else:
                print(f"  [ERROR] Incorrect mimetype: {mimetype}")
            
            # Check file size
            file_size = os.path.getsize(epub_path)
            print(f"  [OK] File size: {file_size:,} bytes ({file_size/1024:.1f} KB)")
            
            return True
            
    except Exception as e:
        print(f"  [ERROR] Error: {e}")
        return False

def main():
    """Main verification function"""
    if len(sys.argv) > 1:
        base_path = sys.argv[1]
    else:
        base_path = r"C:\Users\Stuart\Desktop\Projects\snowflake\artifacts\code_of_deception_20250821_212841\COMPLETE_MANUSCRIPT"
    
    docx_path = f"{base_path}.docx"
    epub_path = f"{base_path}.epub"
    
    print("=== MANUSCRIPT EXPORT VERIFICATION ===")
    print(f"Base path: {base_path}")
    
    # Verify DOCX
    docx_ok = False
    if os.path.exists(docx_path):
        docx_ok = verify_docx(docx_path)
    else:
        print(f"\n[ERROR] DOCX file not found: {docx_path}")
    
    # Verify EPUB
    epub_ok = False
    if os.path.exists(epub_path):
        epub_ok = verify_epub(epub_path)
    else:
        print(f"\n[ERROR] EPUB file not found: {epub_path}")
    
    # Summary
    print(f"\n=== VERIFICATION SUMMARY ===")
    print(f"DOCX: {'PASSED' if docx_ok else 'FAILED'}")
    print(f"EPUB: {'PASSED' if epub_ok else 'FAILED'}")
    
    if docx_ok and epub_ok:
        print("\nSUCCESS: All exports verified successfully!")
        print("The novel 'Code of Deception' is ready for distribution.")
    else:
        print("\nWARNING: Some exports have issues. Please check the details above.")

if __name__ == "__main__":
    main()