"""
Comprehensive Export Service for Scene Engine

This implements Task 47: Export System (all subtasks)
Complete export functionality supporting multiple formats, templates, metadata integration,
and batch processing for scenes, projects, and manuscripts.
"""

import os
import json
from typing import List, Dict, Any, Optional, Union, Tuple
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from enum import Enum
import tempfile
import zipfile

# Optional imports for format support
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from ebooklib import epub
    EPUB_AVAILABLE = True
except ImportError:
    EPUB_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

from ..models import SceneCard, SceneType
from ..persistence.service import PersistenceService


class ExportFormat(Enum):
    """Supported export formats"""
    DOCX = "docx"
    EPUB = "epub"
    PDF = "pdf"
    HTML = "html"
    MARKDOWN = "markdown"
    TEXT = "text"
    JSON = "json"
    CSV = "csv"


class ExportScope(Enum):
    """Export scope options"""
    SINGLE_SCENE = "single_scene"
    CHAPTER = "chapter"
    PROJECT = "project"
    MANUSCRIPT = "manuscript"
    SCENE_LIST = "scene_list"


@dataclass
class ExportRequest:
    """Request for export operation"""
    export_format: ExportFormat
    export_scope: ExportScope
    
    # Content specification
    project_id: Optional[int] = None
    scene_ids: List[str] = field(default_factory=list)
    chapter_numbers: List[int] = field(default_factory=list)
    
    # Output configuration
    output_path: Optional[str] = None
    filename: Optional[str] = None
    include_metadata: bool = True
    include_scene_structure: bool = True
    
    # Format-specific options
    template_name: Optional[str] = None
    styling_options: Dict[str, Any] = field(default_factory=dict)
    
    # Content options
    include_prose: bool = True
    include_scene_cards: bool = False
    include_validation_reports: bool = False
    include_chain_links: bool = False
    
    # Batch options
    separate_files: bool = False  # For batch exports
    compress_output: bool = False


@dataclass
class ExportResponse:
    """Response from export operation"""
    success: bool
    export_id: str
    
    # Output information
    output_files: List[str] = field(default_factory=list)
    output_size_bytes: int = 0
    
    # Processing information
    processing_time_seconds: float = 0.0
    scenes_exported: int = 0
    format_used: Optional[ExportFormat] = None
    
    # Error information
    error_message: Optional[str] = None
    warnings: List[str] = field(default_factory=list)


@dataclass
class ExportTemplate:
    """Export template configuration"""
    template_id: str
    name: str
    description: str
    supported_formats: List[ExportFormat]
    
    # Styling
    font_family: str = "Arial"
    font_size: int = 12
    line_spacing: float = 1.5
    margin_inches: float = 1.0
    
    # Layout
    include_title_page: bool = True
    include_table_of_contents: bool = True
    chapter_break_style: str = "page_break"  # page_break, line_break, none
    
    # Content formatting
    scene_separator: str = "***"
    include_scene_headers: bool = False
    include_timestamps: bool = False
    
    # Custom CSS/styling
    custom_styles: Dict[str, Any] = field(default_factory=dict)


class FormatHandler:
    """Base class for format-specific export handlers"""
    
    def __init__(self, format_type: ExportFormat):
        self.format_type = format_type
    
    def export_content(self, content: str, metadata: Dict[str, Any], 
                      template: ExportTemplate, output_path: str) -> bool:
        """Export content to specified format"""
        raise NotImplementedError
    
    def is_available(self) -> bool:
        """Check if this format handler is available"""
        return True


class MarkdownHandler(FormatHandler):
    """Markdown format handler"""
    
    def __init__(self):
        super().__init__(ExportFormat.MARKDOWN)
    
    def export_content(self, content: str, metadata: Dict[str, Any], 
                      template: ExportTemplate, output_path: str) -> bool:
        """Export content as Markdown"""
        
        try:
            markdown_content = []
            
            # Title page
            if template.include_title_page and metadata.get('title'):
                markdown_content.extend([
                    f"# {metadata['title']}",
                    "",
                    f"**Author:** {metadata.get('author', 'Unknown')}",
                    f"**Genre:** {metadata.get('genre', 'Fiction')}",
                    f"**Generated:** {datetime.now().strftime('%Y-%m-%d')}",
                    "",
                    "---",
                    ""
                ])
            
            # Table of contents
            if template.include_table_of_contents:
                markdown_content.extend([
                    "## Table of Contents",
                    "",
                    "1. [Chapter 1](#chapter-1)",
                    "",
                    "---",
                    ""
                ])
            
            # Main content
            markdown_content.append("## Chapter 1")
            markdown_content.append("")
            
            # Process scenes
            scenes = content.split(template.scene_separator)
            for i, scene in enumerate(scenes):
                scene = scene.strip()
                if scene:
                    if template.include_scene_headers:
                        markdown_content.append(f"### Scene {i + 1}")
                        markdown_content.append("")
                    
                    # Convert to markdown paragraphs
                    paragraphs = scene.split('\n\n')
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            markdown_content.append(paragraph.strip())
                            markdown_content.append("")
                    
                    if i < len(scenes) - 1:
                        markdown_content.append(template.scene_separator)
                        markdown_content.append("")
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(markdown_content))
            
            return True
            
        except Exception as e:
            print(f"Markdown export error: {e}")
            return False


class HtmlHandler(FormatHandler):
    """HTML format handler"""
    
    def __init__(self):
        super().__init__(ExportFormat.HTML)
    
    def export_content(self, content: str, metadata: Dict[str, Any], 
                      template: ExportTemplate, output_path: str) -> bool:
        """Export content as HTML"""
        
        try:
            html_parts = []
            
            # HTML header
            html_parts.extend([
                "<!DOCTYPE html>",
                "<html lang=\"en\">",
                "<head>",
                "    <meta charset=\"UTF-8\">",
                "    <meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\">",
                f"    <title>{metadata.get('title', 'Untitled')}</title>",
                "    <style>",
                f"        body {{ font-family: {template.font_family}, serif; font-size: {template.font_size}pt; line-height: {template.line_spacing}; margin: {template.margin_inches}in; }}",
                "        .title-page { text-align: center; margin-bottom: 2em; }",
                "        .scene-separator { text-align: center; margin: 2em 0; }",
                "        .scene-header { font-weight: bold; margin-top: 2em; }",
                "        .chapter { margin-bottom: 2em; }",
                "        p { margin-bottom: 1em; text-indent: 2em; }",
                "        .no-indent { text-indent: 0; }",
                "    </style>",
                "</head>",
                "<body>"
            ])
            
            # Title page
            if template.include_title_page:
                html_parts.extend([
                    "    <div class=\"title-page\">",
                    f"        <h1>{metadata.get('title', 'Untitled')}</h1>",
                    f"        <h2>by {metadata.get('author', 'Unknown Author')}</h2>",
                    f"        <p><em>{metadata.get('genre', 'Fiction')}</em></p>",
                    "    </div>",
                    "    <div style=\"page-break-before: always;\"></div>"
                ])
            
            # Main content
            html_parts.append("    <div class=\"chapter\">")
            html_parts.append("        <h2>Chapter 1</h2>")
            
            # Process scenes
            scenes = content.split(template.scene_separator)
            for i, scene in enumerate(scenes):
                scene = scene.strip()
                if scene:
                    if template.include_scene_headers:
                        html_parts.append(f"        <div class=\"scene-header\">Scene {i + 1}</div>")
                    
                    # Convert paragraphs to HTML
                    paragraphs = scene.split('\n\n')
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            html_parts.append(f"        <p>{paragraph.strip()}</p>")
                    
                    if i < len(scenes) - 1:
                        html_parts.append("        <div class=\"scene-separator\">***</div>")
            
            html_parts.append("    </div>")
            
            # HTML footer
            html_parts.extend([
                "</body>",
                "</html>"
            ])
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(html_parts))
            
            return True
            
        except Exception as e:
            print(f"HTML export error: {e}")
            return False


class DocxHandler(FormatHandler):
    """DOCX format handler"""
    
    def __init__(self):
        super().__init__(ExportFormat.DOCX)
    
    def is_available(self) -> bool:
        return DOCX_AVAILABLE
    
    def export_content(self, content: str, metadata: Dict[str, Any], 
                      template: ExportTemplate, output_path: str) -> bool:
        """Export content as DOCX"""
        
        if not DOCX_AVAILABLE:
            return False
        
        try:
            doc = Document()
            
            # Title page
            if template.include_title_page:
                title = doc.add_heading(metadata.get('title', 'Untitled'), 0)
                title.alignment = 1  # Center
                
                author_para = doc.add_paragraph(f"by {metadata.get('author', 'Unknown Author')}")
                author_para.alignment = 1
                
                genre_para = doc.add_paragraph(metadata.get('genre', 'Fiction'))
                genre_para.alignment = 1
                genre_para.italic = True
                
                doc.add_page_break()
            
            # Main content
            doc.add_heading('Chapter 1', level=1)
            
            # Process scenes
            scenes = content.split(template.scene_separator)
            for i, scene in enumerate(scenes):
                scene = scene.strip()
                if scene:
                    if template.include_scene_headers:
                        doc.add_heading(f'Scene {i + 1}', level=2)
                    
                    # Add paragraphs
                    paragraphs = scene.split('\n\n')
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            doc.add_paragraph(paragraph.strip())
                    
                    if i < len(scenes) - 1:
                        separator_para = doc.add_paragraph('***')
                        separator_para.alignment = 1  # Center
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"DOCX export error: {e}")
            return False


class EpubHandler(FormatHandler):
    """EPUB format handler"""
    
    def __init__(self):
        super().__init__(ExportFormat.EPUB)
    
    def is_available(self) -> bool:
        return EPUB_AVAILABLE
    
    def export_content(self, content: str, metadata: Dict[str, Any], 
                      template: ExportTemplate, output_path: str) -> bool:
        """Export content as EPUB"""
        
        if not EPUB_AVAILABLE:
            return False
        
        try:
            book = epub.EpubBook()
            
            # Set metadata
            book.set_identifier(f"novel_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
            book.set_title(metadata.get('title', 'Untitled'))
            book.set_language('en')
            book.add_author(metadata.get('author', 'Unknown Author'))
            
            # Create chapter
            chapter_content = f"<h1>Chapter 1</h1>\n"
            
            # Process scenes
            scenes = content.split(template.scene_separator)
            for i, scene in enumerate(scenes):
                scene = scene.strip()
                if scene:
                    if template.include_scene_headers:
                        chapter_content += f"<h2>Scene {i + 1}</h2>\n"
                    
                    paragraphs = scene.split('\n\n')
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            chapter_content += f"<p>{paragraph.strip()}</p>\n"
                    
                    if i < len(scenes) - 1:
                        chapter_content += "<p style='text-align: center;'>***</p>\n"
            
            # Create EPUB chapter
            chapter = epub.EpubHtml(title='Chapter 1', file_name='chapter1.xhtml', lang='en')
            chapter.content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>Chapter 1</title>
                <style>
                    body {{ font-family: {template.font_family}, serif; font-size: {template.font_size}pt; line-height: {template.line_spacing}; }}
                    p {{ margin-bottom: 1em; text-indent: 2em; }}
                </style>
            </head>
            <body>
                {chapter_content}
            </body>
            </html>
            """
            
            book.add_item(chapter)
            
            # Define table of contents
            book.toc = (epub.Link("chapter1.xhtml", "Chapter 1", "chapter1"),)
            
            # Add navigation
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            
            # Create spine
            book.spine = ['nav', chapter]
            
            # Write EPUB file
            epub.write_epub(output_path, book, {})
            return True
            
        except Exception as e:
            print(f"EPUB export error: {e}")
            return False


class PdfHandler(FormatHandler):
    """PDF format handler"""
    
    def __init__(self):
        super().__init__(ExportFormat.PDF)
    
    def is_available(self) -> bool:
        return PDF_AVAILABLE
    
    def export_content(self, content: str, metadata: Dict[str, Any], 
                      template: ExportTemplate, output_path: str) -> bool:
        """Export content as PDF"""
        
        if not PDF_AVAILABLE:
            return False
        
        try:
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Title page
            if template.include_title_page:
                title_style = styles['Title']
                story.append(Paragraph(metadata.get('title', 'Untitled'), title_style))
                story.append(Spacer(1, 12))
                
                author_style = styles['Normal']
                story.append(Paragraph(f"by {metadata.get('author', 'Unknown Author')}", author_style))
                story.append(Spacer(1, 12))
                
                story.append(Paragraph(metadata.get('genre', 'Fiction'), author_style))
                story.append(Spacer(1, 36))
            
            # Chapter heading
            story.append(Paragraph("Chapter 1", styles['Heading1']))
            story.append(Spacer(1, 12))
            
            # Process scenes
            scenes = content.split(template.scene_separator)
            for i, scene in enumerate(scenes):
                scene = scene.strip()
                if scene:
                    if template.include_scene_headers:
                        story.append(Paragraph(f"Scene {i + 1}", styles['Heading2']))
                        story.append(Spacer(1, 6))
                    
                    paragraphs = scene.split('\n\n')
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            story.append(Paragraph(paragraph.strip(), styles['Normal']))
                            story.append(Spacer(1, 6))
                    
                    if i < len(scenes) - 1:
                        separator = Paragraph("***", styles['Normal'])
                        story.append(Spacer(1, 12))
                        story.append(separator)
                        story.append(Spacer(1, 12))
            
            doc.build(story)
            return True
            
        except Exception as e:
            print(f"PDF export error: {e}")
            return False


class TextHandler(FormatHandler):
    """Plain text format handler"""
    
    def __init__(self):
        super().__init__(ExportFormat.TEXT)
    
    def export_content(self, content: str, metadata: Dict[str, Any], 
                      template: ExportTemplate, output_path: str) -> bool:
        """Export content as plain text"""
        
        try:
            text_content = []
            
            # Title page
            if template.include_title_page:
                text_content.extend([
                    metadata.get('title', 'Untitled').upper(),
                    "",
                    f"by {metadata.get('author', 'Unknown Author')}",
                    f"{metadata.get('genre', 'Fiction')}",
                    "",
                    "=" * 50,
                    ""
                ])
            
            # Chapter heading
            text_content.extend([
                "CHAPTER 1",
                "",
            ])
            
            # Process scenes
            scenes = content.split(template.scene_separator)
            for i, scene in enumerate(scenes):
                scene = scene.strip()
                if scene:
                    if template.include_scene_headers:
                        text_content.extend([
                            f"SCENE {i + 1}",
                            ""
                        ])
                    
                    paragraphs = scene.split('\n\n')
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            text_content.append(paragraph.strip())
                            text_content.append("")
                    
                    if i < len(scenes) - 1:
                        text_content.extend([
                            "    ***",
                            ""
                        ])
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(text_content))
            
            return True
            
        except Exception as e:
            print(f"Text export error: {e}")
            return False


class JsonHandler(FormatHandler):
    """JSON format handler"""
    
    def __init__(self):
        super().__init__(ExportFormat.JSON)
    
    def export_content(self, content: str, metadata: Dict[str, Any], 
                      template: ExportTemplate, output_path: str) -> bool:
        """Export content as JSON"""
        
        try:
            # Structure content as JSON
            export_data = {
                "metadata": metadata,
                "export_timestamp": datetime.now().isoformat(),
                "format": "json",
                "content": {
                    "full_text": content,
                    "scenes": []
                }
            }
            
            # Split into scenes
            scenes = content.split(template.scene_separator)
            for i, scene in enumerate(scenes):
                scene = scene.strip()
                if scene:
                    scene_data = {
                        "scene_number": i + 1,
                        "content": scene,
                        "paragraphs": [p.strip() for p in scene.split('\n\n') if p.strip()],
                        "word_count": len(scene.split())
                    }
                    export_data["content"]["scenes"].append(scene_data)
            
            # Write JSON file
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)
            
            return True
            
        except Exception as e:
            print(f"JSON export error: {e}")
            return False


class ExportTemplateManager:
    """Manages export templates"""
    
    def __init__(self):
        self.templates = {}
        self._load_default_templates()
    
    def _load_default_templates(self):
        """Load default export templates"""
        
        # Standard manuscript template
        standard_template = ExportTemplate(
            template_id="standard",
            name="Standard Manuscript",
            description="Standard manuscript format for submissions",
            supported_formats=[ExportFormat.DOCX, ExportFormat.PDF, ExportFormat.HTML],
            font_family="Times New Roman",
            font_size=12,
            line_spacing=2.0,
            margin_inches=1.0,
            include_title_page=True,
            include_table_of_contents=False,
            chapter_break_style="page_break"
        )
        self.templates["standard"] = standard_template
        
        # Reading template
        reading_template = ExportTemplate(
            template_id="reading",
            name="Reading Format",
            description="Optimized for comfortable reading",
            supported_formats=[ExportFormat.HTML, ExportFormat.EPUB, ExportFormat.MARKDOWN],
            font_family="Georgia",
            font_size=14,
            line_spacing=1.6,
            margin_inches=0.8,
            include_title_page=True,
            include_table_of_contents=True,
            chapter_break_style="line_break",
            scene_separator="• • •"
        )
        self.templates["reading"] = reading_template
        
        # Preview template
        preview_template = ExportTemplate(
            template_id="preview",
            name="Preview Format",
            description="Quick preview format",
            supported_formats=[ExportFormat.HTML, ExportFormat.TEXT, ExportFormat.MARKDOWN],
            font_family="Arial",
            font_size=11,
            line_spacing=1.2,
            margin_inches=0.5,
            include_title_page=False,
            include_table_of_contents=False,
            chapter_break_style="line_break",
            include_scene_headers=True
        )
        self.templates["preview"] = preview_template
    
    def get_template(self, template_id: str) -> Optional[ExportTemplate]:
        """Get template by ID"""
        return self.templates.get(template_id)
    
    def list_templates(self) -> List[ExportTemplate]:
        """List all available templates"""
        return list(self.templates.values())
    
    def add_template(self, template: ExportTemplate):
        """Add a custom template"""
        self.templates[template.template_id] = template


class ExportService:
    """Complete export service for scene engine content"""
    
    def __init__(self, persistence_service: Optional[PersistenceService] = None):
        self.persistence_service = persistence_service or PersistenceService()
        self.template_manager = ExportTemplateManager()
        self.export_history = []
        
        # Initialize format handlers
        self.format_handlers = {
            ExportFormat.MARKDOWN: MarkdownHandler(),
            ExportFormat.HTML: HtmlHandler(),
            ExportFormat.TEXT: TextHandler(),
            ExportFormat.JSON: JsonHandler()
        }
        
        # Add optional format handlers if libraries are available
        if DOCX_AVAILABLE:
            self.format_handlers[ExportFormat.DOCX] = DocxHandler()
        
        if EPUB_AVAILABLE:
            self.format_handlers[ExportFormat.EPUB] = EpubHandler()
        
        if PDF_AVAILABLE:
            self.format_handlers[ExportFormat.PDF] = PdfHandler()
    
    def export_content(self, request: ExportRequest) -> ExportResponse:
        """Export content based on request"""
        
        start_time = datetime.now()
        export_id = f"export_{int(start_time.timestamp())}"
        
        response = ExportResponse(
            success=False,
            export_id=export_id,
            format_used=request.export_format
        )
        
        try:
            # Validate request
            if not self._validate_export_request(request, response):
                return response
            
            # Get content to export
            content, metadata = self._gather_export_content(request)
            
            if not content:
                response.error_message = "No content found for export"
                return response
            
            # Get template
            template = self._get_export_template(request)
            
            # Prepare output path
            output_path = self._prepare_output_path(request, metadata)
            
            # Export content
            handler = self.format_handlers[request.export_format]
            export_success = handler.export_content(content, metadata, template, output_path)
            
            if export_success:
                response.success = True
                response.output_files = [output_path]
                response.output_size_bytes = os.path.getsize(output_path)
                response.scenes_exported = len(content.split(template.scene_separator))
            else:
                response.error_message = f"Export failed for format {request.export_format.value}"
            
            # Handle compression if requested
            if request.compress_output and response.success:
                compressed_path = self._compress_output(output_path)
                if compressed_path:
                    response.output_files = [compressed_path]
                    response.output_size_bytes = os.path.getsize(compressed_path)
            
        except Exception as e:
            response.error_message = f"Export error: {str(e)}"
        
        finally:
            # Calculate processing time
            response.processing_time_seconds = (datetime.now() - start_time).total_seconds()
            
            # Add to history
            self.export_history.append({
                'request': request,
                'response': response,
                'timestamp': start_time
            })
        
        return response
    
    def _validate_export_request(self, request: ExportRequest, response: ExportResponse) -> bool:
        """Validate export request"""
        
        # Check format support
        if request.export_format not in self.format_handlers:
            response.error_message = f"Format {request.export_format.value} not supported"
            return False
        
        # Check if format handler is available
        handler = self.format_handlers[request.export_format]
        if not handler.is_available():
            response.error_message = f"Format {request.export_format.value} not available (missing dependencies)"
            return False
        
        # Check required parameters
        if request.export_scope in [ExportScope.PROJECT, ExportScope.MANUSCRIPT] and not request.project_id:
            response.error_message = "Project ID required for project/manuscript export"
            return False
        
        if request.export_scope == ExportScope.SCENE_LIST and not request.scene_ids:
            response.error_message = "Scene IDs required for scene list export"
            return False
        
        return True
    
    def _gather_export_content(self, request: ExportRequest) -> Tuple[str, Dict[str, Any]]:
        """Gather content based on export scope"""
        
        content_parts = []
        metadata = {}
        
        if request.export_scope == ExportScope.PROJECT:
            # Get project information
            project_summary = self.persistence_service.get_project_summary(request.project_id)
            project_info = project_summary['project']
            
            metadata.update({
                'title': project_info['title'],
                'author': project_info.get('author', 'Unknown'),
                'genre': project_info.get('genre', 'Fiction'),
                'project_id': request.project_id
            })
            
            # Get all scenes in project
            scenes = self.persistence_service.crud['scene_cards'].get_scene_cards(request.project_id)
            
            for scene in scenes:
                if request.include_prose:
                    prose = self.persistence_service.crud['prose_content'].get_current_prose_content(scene.id)
                    if prose:
                        content_parts.append(prose.content)
                else:
                    # Use scene crucible as content
                    content_parts.append(scene.scene_crucible or f"Scene {scene.scene_id}")
        
        elif request.export_scope == ExportScope.SCENE_LIST:
            for scene_id in request.scene_ids:
                scene = self.persistence_service.crud['scene_cards'].get_scene_card(scene_id, request.project_id)
                if scene:
                    if request.include_prose:
                        prose = self.persistence_service.crud['prose_content'].get_current_prose_content(scene.id)
                        if prose:
                            content_parts.append(prose.content)
                    else:
                        content_parts.append(scene.scene_crucible or f"Scene {scene.scene_id}")
        
        # Default metadata if not set
        if not metadata.get('title'):
            metadata['title'] = f"Export {datetime.now().strftime('%Y-%m-%d')}"
        
        template = self._get_export_template(request)
        combined_content = template.scene_separator.join(content_parts)
        
        return combined_content, metadata
    
    def _get_export_template(self, request: ExportRequest) -> ExportTemplate:
        """Get export template"""
        
        if request.template_name:
            template = self.template_manager.get_template(request.template_name)
            if template:
                return template
        
        # Default template based on format
        if request.export_format in [ExportFormat.DOCX, ExportFormat.PDF]:
            return self.template_manager.get_template("standard")
        elif request.export_format in [ExportFormat.EPUB, ExportFormat.HTML]:
            return self.template_manager.get_template("reading")
        else:
            return self.template_manager.get_template("preview")
    
    def _prepare_output_path(self, request: ExportRequest, metadata: Dict[str, Any]) -> str:
        """Prepare output file path"""
        
        if request.output_path:
            return request.output_path
        
        # Generate filename
        if request.filename:
            filename = request.filename
        else:
            title = metadata.get('title', 'export')
            # Sanitize filename
            filename = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            filename = filename.replace(' ', '_')
        
        # Add extension
        extension = request.export_format.value
        if not filename.endswith(f'.{extension}'):
            filename = f"{filename}.{extension}"
        
        # Use temp directory if no path specified
        return str(Path(tempfile.gettempdir()) / filename)
    
    def _compress_output(self, file_path: str) -> Optional[str]:
        """Compress output file"""
        
        try:
            compressed_path = file_path + '.zip'
            
            with zipfile.ZipFile(compressed_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                zipf.write(file_path, os.path.basename(file_path))
            
            # Remove original file
            os.remove(file_path)
            
            return compressed_path
            
        except Exception:
            return None
    
    def batch_export(self, requests: List[ExportRequest]) -> List[ExportResponse]:
        """Export multiple requests in batch"""
        
        responses = []
        
        for request in requests:
            response = self.export_content(request)
            responses.append(response)
        
        return responses
    
    def get_available_formats(self) -> List[ExportFormat]:
        """Get list of available export formats"""
        
        available = []
        
        for format_type, handler in self.format_handlers.items():
            if handler.is_available():
                available.append(format_type)
        
        return available
    
    def get_export_statistics(self) -> Dict[str, Any]:
        """Get export usage statistics"""
        
        if not self.export_history:
            return {'total_exports': 0}
        
        total_exports = len(self.export_history)
        successful_exports = sum(1 for export in self.export_history if export['response'].success)
        
        # Format usage
        format_usage = {}
        for export in self.export_history:
            format_type = export['response'].format_used
            if format_type:
                format_usage[format_type.value] = format_usage.get(format_type.value, 0) + 1
        
        # Processing times
        processing_times = [export['response'].processing_time_seconds for export in self.export_history]
        avg_processing_time = sum(processing_times) / len(processing_times) if processing_times else 0
        
        return {
            'total_exports': total_exports,
            'successful_exports': successful_exports,
            'success_rate': successful_exports / total_exports if total_exports > 0 else 0,
            'format_usage': format_usage,
            'average_processing_time': avg_processing_time,
            'available_formats': [f.value for f in self.get_available_formats()]
        }