"""
Manuscript Export Module
Handles exporting manuscripts to DOCX, EPUB, and other formats
"""

import json
from pathlib import Path
from typing import Dict, Any, Optional, TYPE_CHECKING

if TYPE_CHECKING:
    from docx import Document
from datetime import datetime

# Document generation imports
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from ebooklib import epub
    EPUB_AVAILABLE = True
except ImportError:
    EPUB_AVAILABLE = False

class ManuscriptExporter:
    """
    Export manuscripts to various formats (DOCX, EPUB, PDF, etc.)
    """
    
    def __init__(self, project_dir: str = "artifacts"):
        """
        Initialize exporter
        
        Args:
            project_dir: Directory containing project artifacts
        """
        self.project_dir = Path(project_dir)
    
    def export_docx(self, 
                   manuscript: Dict[str, Any],
                   output_path: Optional[Path] = None,
                   project_id: Optional[str] = None) -> Path:
        """
        Export manuscript to DOCX format
        
        Args:
            manuscript: Manuscript data from Step 10
            output_path: Optional output path
            project_id: Project ID for default path
            
        Returns:
            Path to generated DOCX file
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        # Ensure output directory exists
        if output_path:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
        elif project_id:
            export_dir = self.project_dir / project_id
            export_dir.mkdir(parents=True, exist_ok=True)
            output_path = export_dir / "manuscript.docx"
        else:
            self.project_dir.mkdir(parents=True, exist_ok=True)
            output_path = self.project_dir / "manuscript.docx"
        
        # Create document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = manuscript.get('title', 'Untitled Novel')
        doc.core_properties.author = 'Snowflake Method Generator'
        doc.core_properties.created = datetime.now()
        doc.core_properties.keywords = 'novel, fiction, snowflake method'
        
        # Add title page
        self._add_title_page(doc, manuscript)
        
        # Add table of contents
        doc.add_page_break()
        self._add_table_of_contents(doc, manuscript)
        
        # Add chapters and scenes
        for chapter in manuscript.get('chapters', []):
            doc.add_page_break()
            self._add_chapter(doc, chapter)
        
        # Add end matter
        doc.add_page_break()
        self._add_end_matter(doc, manuscript)
        
        # Save document
        if not output_path:
            if project_id:
                project_path = self.project_dir / project_id
                project_path.mkdir(exist_ok=True)
                output_path = project_path / "manuscript.docx"
            else:
                output_path = self.project_dir / "manuscript.docx"
        
        doc.save(str(output_path))
        return output_path
    
    def _add_title_page(self, doc, manuscript: Dict[str, Any]):
        """Add title page to document"""
        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(manuscript.get('title', 'Untitled Novel'))
        title_run.font.size = Pt(36)
        title_run.font.bold = True
        
        # Spacing
        for _ in range(3):
            doc.add_paragraph()
        
        # Word count
        word_para = doc.add_paragraph()
        word_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        word_run = word_para.add_run(f"Word Count: {manuscript.get('total_word_count', 0):,}")
        word_run.font.size = Pt(14)
        
        # Chapter count
        chapter_para = doc.add_paragraph()
        chapter_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chapter_run = chapter_para.add_run(f"Chapters: {manuscript.get('chapter_count', 0)}")
        chapter_run.font.size = Pt(14)
        
        # Bottom matter
        for _ in range(10):
            doc.add_paragraph()
        
        generated_para = doc.add_paragraph()
        generated_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        generated_run = generated_para.add_run("Generated with the Snowflake Method")
        generated_run.font.size = Pt(10)
        generated_run.font.color.rgb = RGBColor(128, 128, 128)
    
    def _add_table_of_contents(self, doc, manuscript: Dict[str, Any]):
        """Add table of contents"""
        toc_heading = doc.add_heading('Table of Contents', level=1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        for chapter in manuscript.get('chapters', []):
            chapter_num = chapter.get('number', 0)
            word_count = chapter.get('word_count', 0)
            
            toc_entry = doc.add_paragraph(f"Chapter {chapter_num}", style='List Bullet')
            toc_entry.add_run(f" ({word_count:,} words)").font.color.rgb = RGBColor(128, 128, 128)
            
            # Add scene list for chapter
            for scene in chapter.get('scenes', []):
                scene_num = scene.get('scene_num', 0)
                pov = scene.get('pov', 'Unknown')
                scene_entry = doc.add_paragraph(f"    Scene {scene_num} - {pov}", style='List Bullet 2')
                
                if scene.get('disaster_anchor'):
                    scene_entry.add_run(f" [{scene['disaster_anchor']}]").font.bold = True
    
    def _add_chapter(self, doc, chapter: Dict[str, Any]):
        """Add a chapter to the document"""
        chapter_num = chapter.get('number', 0)
        
        # Chapter heading
        chapter_heading = doc.add_heading(f'Chapter {chapter_num}', level=1)
        chapter_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add scenes
        for i, scene in enumerate(chapter.get('scenes', [])):
            if i > 0:
                # Scene break
                doc.add_paragraph('* * *', style='Intense Quote')
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            self._add_scene(doc, scene)
    
    def _add_scene(self, doc, scene: Dict[str, Any]):
        """Add a scene to the document"""
        # Scene metadata (optional - can be removed for final version)
        if scene.get('disaster_anchor'):
            metadata = doc.add_paragraph()
            metadata.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            metadata_run = metadata.add_run(f"[{scene['disaster_anchor']} - {scene.get('type', 'Scene')}]")
            metadata_run.font.size = Pt(10)
            metadata_run.font.italic = True
            metadata_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Scene prose
        prose = scene.get('prose', '')
        
        # Split into paragraphs
        for para_text in prose.split('\n\n'):
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.paragraph_format.first_line_indent = Inches(0.5)
                para.paragraph_format.line_spacing = 1.5
    
    def _add_end_matter(self, doc, manuscript: Dict[str, Any]):
        """Add end matter to document"""
        doc.add_heading('The End', level=2)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Statistics
        stats_heading = doc.add_heading('Manuscript Statistics', level=3)
        stats_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        stats = [
            f"Total Words: {manuscript.get('total_word_count', 0):,}",
            f"Chapters: {manuscript.get('chapter_count', 0)}",
            f"Scenes: {manuscript.get('scene_count', 0)}",
            f"Generated: {manuscript.get('metadata', {}).get('created_at', 'Unknown')}"
        ]
        
        for stat in stats:
            stat_para = doc.add_paragraph(stat)
            stat_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            stat_para.runs[0].font.size = Pt(11)
    
    def export_epub(self,
                   manuscript: Dict[str, Any],
                   output_path: Optional[Path] = None,
                   project_id: Optional[str] = None) -> Path:
        """
        Export manuscript to EPUB format
        
        Args:
            manuscript: Manuscript data from Step 10
            output_path: Optional output path
            project_id: Project ID for default path
            
        Returns:
            Path to generated EPUB file
        """
        if not EPUB_AVAILABLE:
            raise ImportError("ebooklib not installed. Run: pip install ebooklib")
        
        # Ensure output directory exists
        if output_path:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
        elif project_id:
            export_dir = self.project_dir / project_id
            export_dir.mkdir(parents=True, exist_ok=True)
            output_path = export_dir / "manuscript.epub"
        else:
            self.project_dir.mkdir(parents=True, exist_ok=True)
            output_path = self.project_dir / "manuscript.epub"
        
        # Create EPUB book
        book = epub.EpubBook()
        
        # Set metadata
        book.set_identifier(f"snowflake_{project_id or 'novel'}")
        book.set_title(manuscript.get('title', 'Untitled Novel'))
        book.set_language('en')
        book.add_author('Snowflake Method Generator')
        
        # Create chapters
        epub_chapters = []
        spine = ['nav']
        
        for chapter_data in manuscript.get('chapters', []):
            chapter_num = chapter_data.get('number', 0)
            
            # Create chapter
            chapter = epub.EpubHtml(
                title=f'Chapter {chapter_num}',
                file_name=f'chap_{chapter_num:02d}.xhtml',
                lang='en'
            )
            
            # Build chapter content
            content = f'<h1>Chapter {chapter_num}</h1>\n'
            
            for i, scene in enumerate(chapter_data.get('scenes', [])):
                if i > 0:
                    content += '<p style="text-align: center;">* * *</p>\n'
                
                # Add scene prose
                prose = scene.get('prose', '')
                for para in prose.split('\n\n'):
                    if para.strip():
                        content += f'<p>{para.strip()}</p>\n'
            
            chapter.content = content
            
            # Add to book
            book.add_item(chapter)
            epub_chapters.append(chapter)
            spine.append(chapter)
        
        # Add navigation
        book.toc = tuple(epub_chapters)
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Set spine
        book.spine = spine
        
        # Add CSS
        style = '''
        @namespace epub "http://www.idpf.org/2007/ops";
        body {
            font-family: Georgia, serif;
            margin: 5%;
            text-align: justify;
        }
        h1 {
            text-align: center;
            margin-bottom: 2em;
        }
        p {
            text-indent: 1.5em;
            margin: 0;
        }
        p:first-of-type {
            text-indent: 0;
        }
        '''
        
        nav_css = epub.EpubItem(
            uid="style_nav",
            file_name="style/nav.css",
            media_type="text/css",
            content=style
        )
        book.add_item(nav_css)
        
        # Save EPUB
        if not output_path:
            if project_id:
                project_path = self.project_dir / project_id
                project_path.mkdir(exist_ok=True)
                output_path = project_path / "manuscript.epub"
            else:
                output_path = self.project_dir / "manuscript.epub"
        
        epub.write_epub(str(output_path), book, {})
        return output_path
    
    def export_markdown(self,
                       manuscript: Dict[str, Any],
                       output_path: Optional[Path] = None,
                       project_id: Optional[str] = None) -> Path:
        """
        Export manuscript to Markdown format
        
        Args:
            manuscript: Manuscript data from Step 10
            output_path: Optional output path
            project_id: Project ID for default path
            
        Returns:
            Path to generated Markdown file
        """
        # Ensure output directory exists
        if output_path:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
        elif project_id:
            export_dir = self.project_dir / project_id
            export_dir.mkdir(parents=True, exist_ok=True)
            output_path = export_dir / "manuscript.md"
        else:
            self.project_dir.mkdir(parents=True, exist_ok=True)
            output_path = self.project_dir / "manuscript.md"
        
        # Extract manuscript data
        ms_data = manuscript.get("manuscript", manuscript)
        
        # Write markdown
        with open(output_path, "w", encoding="utf-8") as f:
            # Title and metadata
            f.write(f"# {manuscript.get('title', 'Manuscript')}\n\n")
            
            metadata = manuscript.get('metadata', {})
            if metadata.get('author'):
                f.write(f"**Author:** {metadata['author']}\n\n")
            
            f.write(f"**Total Words:** {ms_data.get('total_word_count', 0):,}\n\n")
            f.write("---\n\n")
            
            # Chapters and scenes
            for chapter in ms_data.get('chapters', []):
                chapter_title = chapter.get('title', f"Chapter {chapter.get('number', '?')}")
                f.write(f"## {chapter_title}\n\n")
                
                for scene in chapter.get('scenes', []):
                    f.write(f"### Scene {scene.get('scene_number', '?')}\n\n")
                    f.write(f"*POV: {scene.get('pov', 'Unknown')} | Type: {scene.get('type', 'Unknown')}*\n\n")
                    f.write(scene.get('prose', '[No prose generated]'))
                    f.write("\n\n---\n\n")
        
        return output_path
    
    def export_all_formats(self,
                          manuscript: Dict[str, Any],
                          project_id: str) -> Dict[str, Path]:
        """
        Export manuscript to all available formats
        
        Args:
            manuscript: Manuscript data from Step 10
            project_id: Project ID
            
        Returns:
            Dictionary of format names to file paths
        """
        exports = {}
        
        # Always export markdown (already done in Step 10)
        project_path = self.project_dir / project_id
        exports['markdown'] = project_path / "manuscript.md"
        exports['text'] = project_path / "manuscript.txt"
        exports['json'] = project_path / "step_10_manuscript.json"
        
        # Export DOCX if available
        if DOCX_AVAILABLE:
            try:
                exports['docx'] = self.export_docx(manuscript, project_id=project_id)
                print(f"✓ Exported DOCX: {exports['docx']}")
            except Exception as e:
                print(f"✗ DOCX export failed: {e}")
        
        # Export EPUB if available
        if EPUB_AVAILABLE:
            try:
                exports['epub'] = self.export_epub(manuscript, project_id=project_id)
                print(f"✓ Exported EPUB: {exports['epub']}")
            except Exception as e:
                print(f"✗ EPUB export failed: {e}")
        
        return exports